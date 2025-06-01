"""
Script para generar un documento Word de ejemplo usando python-docx.
Este script crea un archivo DOCX simple que puede ser usado como ejemplo.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Obtener la ruta del directorio actual
current_dir = os.path.dirname(os.path.abspath(__file__))

# Crear la carpeta de ejemplos si no existe
example_dir = os.path.join(current_dir, 'static', 'examples')
os.makedirs(example_dir, exist_ok=True)

# Crear un nuevo documento
document = Document()

# Establecer título
title = document.add_heading('Documento de Ejemplo', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Añadir subtítulo
subtitle = document.add_paragraph('Para probar la conversión de Word a PDF')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.runs[0]
subtitle_run.font.italic = True
subtitle_run.font.size = Pt(14)
subtitle_run.font.color.rgb = RGBColor(0x44, 0x44, 0xBB)

# Añadir una descripción
document.add_paragraph('Este es un documento de ejemplo para probar la funcionalidad de conversión de Word a PDF. El conversor utiliza LibreOffice para realizar la transformación de archivos DOCX a formato PDF de alta calidad.')

# Añadir una sección con características
document.add_heading('Características principales', level=1)
features = [
    'Conversión rápida y precisa',
    'Mantiene el formato original',
    'Soporta documentos complejos',
    'Interfaz web moderna y fácil de usar'
]

for feature in features:
    p = document.add_paragraph()
    p.add_run('• ').bold = True
    p.add_run(feature)

# Añadir una imagen (comentado porque puede que no exista)
# document.add_heading('Logo', level=1)
# document.add_picture('path/to/logo.png', width=Inches(2.0))

# Añadir una tabla
document.add_heading('Formatos soportados', level=1)
table = document.add_table(rows=1, cols=3)
table.style = 'Table Grid'

# Establecer encabezados
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Formato'
hdr_cells[1].text = 'Extensión'
hdr_cells[2].text = 'Notas'

# Añadir datos
row_data = [
    ('Word', '.docx', 'Formato recomendado'),
    ('Word (antiguo)', '.doc', 'Soporte limitado'),
    ('OpenDocument', '.odt', 'Compatible con LibreOffice')
]

for format_name, extension, notes in row_data:
    row_cells = table.add_row().cells
    row_cells[0].text = format_name
    row_cells[1].text = extension
    row_cells[2].text = notes

# Añadir una nota al pie
document.add_paragraph().add_run().add_break()
footer = document.add_paragraph('Documento generado automáticamente como ejemplo.')
footer.runs[0].font.size = Pt(8)
footer.runs[0].font.italic = True

# Guardar el documento
output_path = os.path.join(example_dir, 'documento_ejemplo.docx')
document.save(output_path)
print(f"Documento de ejemplo creado en: {output_path}")
